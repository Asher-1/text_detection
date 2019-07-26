# -*- coding: utf-8 -*-
import os
import time
from docx import Document
from docx.shared import Inches
from sgmllib import SGMLParser
from urllib.request import urlopen

# import pdfkit
#
# pdfkit.from_url('http://baidu.com', 'url.pdf')
#
# # pdfkit.from_file('test.html','out.pdf')
# with open("pred.txt", "r") as f:
#     text = f.read()
# pdfkit.from_string(text, 'file.pdf')


##获取要解析的url
class GetUrl(SGMLParser):
    def __init__(self):
        SGMLParser.__init__(self)
        self.start = False
        self.urlArr = []

    def start_div(self, attr):
        for name, value in attr:
            if value == "ChairmanCont Bureau":  # 页面js中的固定值
                self.start = True

    def end_div(self):
        self.start = False

    def start_a(self, attr):
        if self.start:
            for name, value in attr:
                self.urlArr.append(value)

    def getUrlArr(self):
        return self.urlArr


##解析上面获取的url，获取有用数据
class getManInfo(SGMLParser):
    def __init__(self):
        SGMLParser.__init__(self)
        self.start = False
        self.p = False
        self.dl = False
        self.manInfo = []
        self.subInfo = []

    def start_div(self, attr):
        for name, value in attr:
            if value == "SpeakerInfo":  # 页面js中的固定值
                self.start = True

    def end_div(self):
        self.start = False

    def start_p(self, attr):
        if self.dl:
            self.p = True

    def end_p(self):
        self.p = False

    def start_img(self, attr):
        if self.dl:
            for name, value in attr:
                self.subInfo.append(value)

    def handle_data(self, data):
        if self.p:
            self.subInfo.append(data.decode('utf-8'))

    def start_dl(self, attr):
        if self.start:
            self.dl = True

    def end_dl(self):
        self.manInfo.append(self.subInfo)
        self.subInfo = []
        self.dl = False

    def getManInfo(self):
        return self.manInfo


# urlSource = "http://www.baidu.com"
urlSource = "https://www.csdn.net/gather_28/MtjaEgwsNTA5LWJsb2cO0O0O.html"
sourceData = urlopen(urlSource).read().decode('utf-8')

startTime = time.clock()
##get urls
getUrl = GetUrl()
getUrl.feed(sourceData)
urlArr = getUrl.getUrlArr()
getUrl.close()
print("get url use:" + str((time.clock() - startTime)))
startTime = time.clock()

##get maninfos
manInfos = getManInfo()
for url in urlArr:  # one url one person
    data = urlopen(url).read()
    manInfos.feed(data)
infos = manInfos.getManInfo()
print(infos)
manInfos.close()
print("get maninfos use:" + str((time.clock() - startTime)))
startTime = time.clock()

# word
saveFile = os.getcwd() + "/temp.docx"
doc = Document()
##word title
doc.add_heading("HEAD", level=0)
p = doc.add_paragraph("HEADCONTENT:")

##write info
for infoArr in infos:
    i = 0
    for info in infoArr:
        if i == 0:  ##img url
            arr1 = info.split('.')
            suffix = arr1[len(arr1) - 1]
            arr2 = info.split('/')
            preffix = arr2[len(arr2) - 2]
            imgFile = os.getcwd() + "\\imgs\\" + preffix + "." + suffix
            if not os.path.exists(os.getcwd() + "\\imgs"):
                os.mkdir(os.getcwd() + "\\imgs")
            imgData = urlopen(info).read()

            try:
                f = open(imgFile, 'wb')
                f.write(imgData)
                f.close()
                doc.add_picture(imgFile, width=Inches(1.25))
                os.remove(imgFile)
            except Exception as err:
                print(err)

        elif i == 1:
            doc.add_heading(info + ":", level=1)
        else:
            doc.add_paragraph(info, style='ListBullet')
        i = i + 1

doc.save(saveFile)
print("word use:" + str((time.clock() - startTime)))
